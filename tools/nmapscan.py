"""Nmap port scanning tool"""
import nmap
from docx import Document

def print_and_add_paragraph(document, message):
    """Skriv ut och lägg till meddelande som ett stycke i ett Word-dokument."""
    print(message)
    document.add_paragraph(message)

def scan_host(nm, host, document):
    """Skanna en specifik värd och lägg till resultaten i ett Word-dokument."""
    host_info = f'Host: {host} ({nm[host].hostname()})'
    print_and_add_paragraph(document, host_info)

    state_info = f'State: {nm[host].state()}'
    print_and_add_paragraph(document, state_info)

    for proto in nm[host].all_protocols():
        proto_info = '----------\n' + f'Protocol: {proto}'
        print_and_add_paragraph(document, proto_info)

        lport = sorted(nm[host][proto].keys())
        for port in lport:
            if nm[host][proto][port]['state'] == 'open':
                port_info = f'Port: {port}\tState: {nm[host][proto][port]["state"]}'
                print_and_add_paragraph(document, port_info)
                if 'version' in nm[host][proto][port]:
                    version_info = f'Version: {nm[host][proto][port]["version"]}'
                    print_and_add_paragraph(document, version_info)

def get_nmapscan(ip_address, output_file='nmap_scan_results.docx'):
    """Kör en portscanning via nmap och sparar resultaten i ett Word-dokument."""
    nm = nmap.PortScanner()
    doc = Document()

    try:
        nm.scan(ip_address, arguments='-sV -sC -p- --open -v')
    except nmap.PortScannerError as e:
        print(f"Skanning misslyckades: {e}")
        return
    except Exception as e:  # Behåller generell exceptionhantering som en sista utväg
        print(f"Okänt fel: {e}")
        return

    hosts_up = nm.all_hosts()  # Hämta listan över upptäckta värdar

    if hosts_up:
        for host in hosts_up:
            scan_host(nm, host, doc)
    else:
        print_and_add_paragraph(doc, "Inga värdar hittades.")

    doc.save(output_file)

# Ersätt 'IP_ADDRESS' med den faktiska IP-adressen du vill skanna
# get_nmapscan('IP_ADDRESS')
